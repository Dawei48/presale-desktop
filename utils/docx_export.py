"""
Word 导出 - 仅导出未发货订单
格式: 产品图片 | 产品名 | 价格 | 预售数量
按产品合并，压缩图片
"""
import os
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PIL import Image as PILImage


def compress_image(img_path, max_size=(300, 300), quality=60):
    """压缩图片并返回字节流（支持本地路径和URL）"""
    if not img_path:
        return None
    try:
        if img_path.startswith("http"):
            import urllib.request
            with urllib.request.urlopen(img_path, timeout=10) as resp:
                img = PILImage.open(resp)
        else:
            img = PILImage.open(img_path)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        img.thumbnail(max_size, PILImage.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=quality, optimize=True)
        buf.seek(0)
        return buf
    except Exception:
        return None


def export_brand_orders_docx(brand_name, orders, filepath, title=None):
    """导出未发货订单 - 按产品合并"""
    # 只取未发货订单
    pending = [o for o in orders if o.get("status") == "pending"]
    if not pending:
        pending = orders  # 如果没有筛选到，就用全部（兜底）

    if title is None:
        title = f"{brand_name} - 预售订单"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # Logo + 标题
    from config import BASE_DIR
    from config import _find_resource
    logo_path = _find_resource("assets/logo.png")
    if os.path.exists(logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_path, width=Cm(2.5))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"生成日期：{datetime.now().strftime('%Y-%m-%d')}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()

    # 按产品合并: 同一产品多笔订单 → 一行
    product_map = {}
    for o in pending:
        pname = o.get("product_name", "未知")
        if pname not in product_map:
            product_map[pname] = {
                "price": o.get("product_price", 0),
                "image_path": o.get("product_image_path", ""),
                "total_qty": 0,
                "orders": [],
                "dates": [],
            }
        product_map[pname]["total_qty"] += o.get("quantity", 0)
        product_map[pname]["orders"].append(o)
        date_str = o.get("created_at", "")[:10]
        if date_str and date_str not in product_map[pname]["dates"]:
            product_map[pname]["dates"].append(date_str)

    if not product_map:
        doc.add_paragraph("暂无未发货订单")
        doc.save(filepath)
        return filepath

    # 表格: 图片 | 产品名称 | 价格 | 预售数量 | 预售日期
    headers = ["产品图片", "产品名称", "价格", "预售数量", "预售日期"]
    table = doc.add_table(rows=1 + len(product_map), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # 设置列宽
    for row in table.rows:
        row.cells[0].width = Cm(2.5)
        row.cells[1].width = Cm(4)
        row.cells[2].width = Cm(2)
        row.cells[3].width = Cm(2)
        row.cells[4].width = Cm(5)

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for par in cell.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="0F172A"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # 数据行
    for ri, (pname, pdata) in enumerate(product_map.items()):
        row = table.rows[ri + 1]
        row.height = Cm(2.5)

        # 图片列
        cell_img = row.cells[0]
        cell_img.paragraphs[0].clear()
        img_path = pdata.get("image_path", "")
        has_img = img_path and (img_path.startswith("http") or os.path.exists(img_path))
        if has_img:
            buf = compress_image(img_path, max_size=(200, 200), quality=50)
            if buf:
                p = cell_img.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(buf, width=Cm(2.2))

        # 产品名
        cell_name = row.cells[1]
        cell_name.text = pname
        for par in cell_name.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in par.runs:
                run.font.size = Pt(11)
                run.font.bold = True

        # 价格
        cell_price = row.cells[2]
        cell_price.text = f"¥{pdata['price']:,.0f}"
        for par in cell_price.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.size = Pt(11)

        # 预售数量
        cell_qty = row.cells[3]
        cell_qty.text = str(pdata["total_qty"])
        for par in cell_qty.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in par.runs:
                run.font.size = Pt(11)
                run.font.bold = True

        # 预售日期（所有订单日期，逗号分隔）
        cell_dates = row.cells[4]
        dates = pdata.get("dates", [])
        cell_dates.text = "、".join(dates) if dates else "-"
        for par in cell_dates.paragraphs:
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in par.runs:
                run.font.size = Pt(9)

        # 斑马纹
        if ri % 2 == 0:
            for ci in range(len(headers)):
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
                row.cells[ci]._tc.get_or_add_tcPr().append(shading)

    # 页脚
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— 放心预 —")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    os.makedirs(os.path.dirname(filepath) or ".", exist_ok=True)
    doc.save(filepath)
    return filepath
